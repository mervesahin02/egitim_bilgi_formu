from docx import Document
from pathlib import Path
import os

def generate_word(data):
    """
    JSON verisini kullanarak bir Word (DOCX) dosyası oluşturur.
    """
    try:
        doc = Document()

        # Başlık ekle
        doc.add_heading('Eğitim Bilgi Formu', level=1)

        # Verileri belgeye ekle
        for key, value in data.items():
            doc.add_paragraph(f"{key}: {value}")

        # Word dosyasını kaydet
        output_path = Path("output") / "egitim_bilgileri.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)  # Klasör yoksa oluştur
        doc.save(output_path)

        print(f"✅ Word dosyası başarıyla oluşturuldu: {output_path}")
        return output_path  # Artık Path objesi döndürüyorsun
    
    except Exception as e:
        print(f"❌ Word dosyası oluşturulurken bir hata oluştu: {e}")
        raise Exception("Word oluşturma sırasında bir hata oluştu.")
