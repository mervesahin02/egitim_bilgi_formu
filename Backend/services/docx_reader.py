from docx import Document
from io import BytesIO
from thefuzz import fuzz, process  
import traceback  
from Backend.services.data_validator import find_best_match

def read_docx(file):
    """DOCX dosyasını okur ve JSON formatında döndürür."""
    try:
        doc = Document(BytesIO(file))  

        data = {
            "id": "",  # ✅ ID Bilgisi eklendi
            "egitim_adi": "",
            "egitmen_adi": "",
            "egitim_suresi": "",
            "hedef_kitle": "",
            "egitim_ozeti": "",
            "kaynak_dokumanlar": ""
        }

        print("\n🔍 DOCX Dosyası Okunuyor...")
        print("📄 Tüm Paragraflar:", [p.text for p in doc.paragraphs])

        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    key_text = row.cells[0].text.strip()
                    value_text = row.cells[1].text.strip()

                    if key_text:
                        best_match = find_best_match(key_text, data.keys())
                        if best_match:
                            data[best_match] = value_text
                            print(f"✅ EŞLEŞTİ: {key_text} -> {best_match} : {value_text}")

        # ID eksikse otomatik oluştur
        if not data["id"]:
            import uuid
            data["id"] = str(uuid.uuid4())[:8]  # ✅ UUID ile kısa ID oluştur

        print("\n📊 Okunan Veri:", data)
        return data

    except Exception as e:
        error_message = f"DOCX işleme hatası: {str(e)}\n{traceback.format_exc()}"
        print(error_message)
        raise Exception(error_message)
